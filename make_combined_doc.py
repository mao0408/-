from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("聊小智_运行原理制作步骤与优化方向_修正版.docx")


BLUE = RGBColor(31, 78, 121)
GREEN = RGBColor(0, 176, 80)
GRAY = RGBColor(89, 89, 89)
LIGHT_GREEN = "EAF7EF"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F7"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_font(run, size=11, bold=False, color=None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_p(doc: Document, text: str = "", style: str | None = None, bold=False, color=None, size=11):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, color=color)
    return p


def add_title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("聊小智运行原理、制作步骤与后续优化方向")
    set_font(r, size=22, bold=True, color=BLUE)
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Windows 当前聊天窗口 OCR + 大模型回复助手学习文档")
    set_font(r, size=11, color=GRAY)
    p.paragraph_format.space_after = Pt(18)

    add_callout(
        doc,
        "文档定位",
        "这份文档用于让其他开发者或学习者理解该程序的设计逻辑、关键模块、实现步骤、交付方式和后续优化路线。它不是营销说明，而是偏工程复盘和制作指南。",
        LIGHT_BLUE,
    )


def add_h1(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_font(r, size=16, bold=True, color=BLUE)
    return p


def add_h2(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_font(r, size=13, bold=True, color=BLUE)
    return p


def add_body(doc: Document, text: str):
    p = add_p(doc, text)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Cm(0.65)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        set_font(r, size=10.5)


def add_numbers(doc: Document, items: list[str]) -> None:
    for index, item in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.72)
        p.paragraph_format.first_line_indent = Cm(-0.72)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(f"{index}. {item}")
        set_font(r, size=10.5)


def add_callout(doc: Document, title: str, text: str, fill: str = LIGHT_GRAY) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Cm(16.0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_font(r, size=10.5, bold=True, color=BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    set_font(r2, size=10)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for i, width in enumerate(widths_cm):
        table.columns[i].width = Cm(width)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, LIGHT_BLUE)
        set_cell_margins(cell)
        r = cell.paragraphs[0].add_run(h)
        set_font(r, size=10, bold=True, color=BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_margins(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            r = cells[i].paragraphs[0].add_run(value)
            set_font(r, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def setup_doc() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)

    return doc


def build() -> None:
    doc = setup_doc()
    add_title(doc)

    add_h1(doc, "一、项目目标")
    add_body(
        doc,
        "聊小智的目标是在不调用微信、QQ、企业微信等官方聊天接口的前提下，通过桌面截图、OCR、当前窗口识别、大模型生成和本地话术库检索，为用户提供当前聊天窗口的候选回复。用户可以手动识别、编辑、发送，也可以在风险可控的前提下开启自动识别或全托管。",
    )
    add_bullets(
        doc,
        [
            "核心定位：可视化伴随工具，不侵入聊天软件数据库，也不需要聊天平台 API。",
            "第一优先级：识别当前正在处理的聊天窗口，围绕最新未回复内容生成回复。",
            "交互方式：显示上下文、候选回复、手动输入框、发送按钮和回复来源状态。",
            "数据策略：用户确认发送或训练模式采集的内容进入本地话术库；全托管发送不写入记忆。",
        ],
    )

    add_h1(doc, "二、整体运行流程")
    add_numbers(
        doc,
        [
            "发现当前聊天窗口：读取 Windows 前台窗口、窗口标题、进程名和应用特征，判断当前目标是微信、QQ、企业微信或其他可识别聊天窗口。",
            "定位聊天区域：根据窗口结构、聊天列表宽度、顶部标题栏、底部输入框和消息区边界，动态截取真正的聊天内容区域。",
            "截图与 OCR：对聊天区截图，调用本地 OCR 识别文字，并结合气泡位置、头像距离、颜色和布局判断消息来自对方还是自己。",
            "上下文清洗：过滤时间戳、按钮、语音时长、图片噪声、无意义乱码和重复行，只保留可用于回复的对话上下文。",
            "可见内容理解：遇到图片、卡片、公众号标题时，只使用屏幕上可靠可见的信息；不把视觉模型脑补内容当作聊天原文。",
            "生成候选回复：根据回复来源设置，调用大模型、话术库或混合模式生成 3 条候选。",
            "用户确认发送：用户可改、复制或发送；成功发送后，手动确认的回复进入长期记忆。",
            "全托管循环：只在最新一轮对方未回复消息出现时生成并发送一条；发送后等待新对方消息，不连续自我回复。",
        ],
    )

    add_h1(doc, "三、核心模块说明")
    add_table(
        doc,
        ["模块", "职责", "关键点"],
        [
            ["窗口识别", "识别当前聊天软件和联系人", "依赖 Windows 窗口标题、进程名、尺寸和界面特征，避免误把 Codex、浏览器等窗口当聊天窗口。"],
            ["区域定位", "确定聊天内容截图范围", "不能写死坐标，需要随窗口全屏、半屏、拖拽变化动态计算。"],
            ["OCR 清洗", "读取聊天文字并过滤噪声", "识别结果要去掉时间、按钮、语音秒数、图片参数和明显乱码。"],
            ["角色判断", "区分对方消息和我方消息", "不能只看左右比例，长文本会跨中线；需要结合头像距离、气泡边界、颜色和连续行合并。"],
            ["回复引擎", "生成候选回复", "大模型模式不掺入话术库；话术库模式优先检索本地相似样例；混合模式先检索再让模型理解上下文。"],
            ["话术库", "保存长期记忆和导入语料", "导入 txt、docx、xlsx、xls、csv、json、html、zip 后，解析为上下文和回复样例。"],
            ["发送器", "把回复发送到当前聊天窗口", "通过窗口聚焦、剪贴板粘贴和 Enter 发送；发送前校验目标窗口没有切换。"],
            ["设置与打包", "配置 API、视觉模型、话术库导入和便携交付", "便携版不带 API Key 和训练数据，客户首次打开后自己配置。"],
        ],
        [2.3, 4.0, 9.2],
    )

    add_h1(doc, "四、为什么不需要聊天软件官方 API")
    add_body(
        doc,
        "程序本质上是模拟人类看屏幕和输入文字的过程：它通过 Windows 截图读取聊天窗口可见内容，用 OCR 识别文字，再通过键鼠自动化把最终回复粘贴到当前输入框。因此它不需要读取微信数据库，也不需要调用微信官方接口。",
    )
    add_table(
        doc,
        ["能力", "实现方式", "边界"],
        [
            ["读取聊天", "截取屏幕可见聊天区域并 OCR", "只能识别当前屏幕能看到的内容，滚动历史需要额外处理。"],
            ["判断对象", "识别当前窗口标题和联系人名称", "窗口遮挡、误识别标题或多窗口叠放时需要保护逻辑。"],
            ["发送消息", "聚焦目标窗口，粘贴文本并回车", "必须确认目标窗口未切换，否则可能误发。"],
            ["学习风格", "保存用户确认发送的上下文和回复", "不是微调大模型，而是本地长期记忆和相似检索增强。"],
        ],
        [2.5, 5.4, 7.6],
    )

    add_h1(doc, "五、回复来源与生成逻辑")
    add_h2(doc, "1. 大模型模式")
    add_body(doc, "只把当前上下文和目标未回复轮次发给所选文本模型，例如 DeepSeek 或智谱。该模式不应该混入本地话术库内容，适合没有训练数据的新客户。")
    add_h2(doc, "2. 话术库模式")
    add_body(doc, "优先从本地长期记忆或导入语料中检索相似上下文，并直接给出更接近历史话术的候选。如果没有足够相似样例，再用大模型兜底。")
    add_h2(doc, "3. 混合模式")
    add_body(doc, "先检索本地相似话术，再把相似样例、当前上下文和最新未回复问题一起交给大模型。该模式适合已有行业话术，希望既稳定又能理解新语境。")
    add_callout(
        doc,
        "重要边界",
        "设置里导入的训练语料并不是对 DeepSeek 或智谱做真实微调，而是生成本地话术库和相似检索样例。它属于 RAG/长期记忆增强，不是模型权重训练。",
        LIGHT_GREEN,
    )

    add_h1(doc, "六、话术库导入与本地记忆")
    add_body(
        doc,
        "话术库的目标是让客户上传自己的聊天记录、客服话术表或历史问答，程序在本地解析成“对方上下文 -> 我方回复”的样例。后续识别到相似问题时，程序可以优先参考这些真实回复。",
    )
    add_bullets(
        doc,
        [
            "支持文件：txt、md、csv、tsv、json、docx、xlsx、xls、html、zip，以及包含这些文件的文件夹。",
            "支持格式：A/B 对话、问题/回复表、客服宝类话术表、两人聊天记录、带时间和发送人的聊天导出文本。",
            "去重方式：根据上下文、回复内容和来源生成哈希，避免重复导入同一批数据。",
            "保存位置：源码版通常在用户 AppData 下；便携版应保存在程序同级 user_data 目录。",
            "记忆写入规则：手动发送、候选发送、手动输入发送、训练模式观察到的我方回复可以写入；全托管无论成功与否都不写入。",
        ],
    )

    add_h1(doc, "七、全托管逻辑")
    add_body(
        doc,
        "全托管不是自动发送第一条候选，而是先识别当前未回复轮次，再生成候选，再按审核规则选择一条发送。核心原则是一轮只回复一次，回复后必须等待对方新增消息。",
    )
    add_table(
        doc,
        ["判断场景", "处理方式"],
        [
            ["最新有效消息是我方", "不生成、不发送，继续监听。"],
            ["最近一条我方消息之后有连续对方消息", "把这一组作为当前未回复轮次，统一生成一条回复。"],
            ["同一轮对方消息已处理", "不重复发送，避免连续自动回复。"],
            ["发送后 OCR 仍看到同一轮消息", "仍然不回复；只有对方新增消息才进入下一轮。"],
            ["目标窗口切换或联系人变化", "发送前拦截，避免发到错误窗口。"],
        ],
        [5.0, 10.0],
    )

    add_h1(doc, "八、图片、卡片与可见内容理解")
    add_body(
        doc,
        "图片和卡片是 OCR 方案的难点。程序可以识别屏幕上能看见的标题、按钮、图片中的文字或文件名，但不能假装打开链接，也不能把视觉模型的主观脑补当作聊天内容。",
    )
    add_bullets(
        doc,
        [
            "表情包：默认不回复，除非用户手动输入。",
            "普通图片：优先 OCR 图片中真实可见文字；无法识别时显示为图片，不生成虚构上下文。",
            "公众号/卡片：只使用可见标题、摘要和卡片类型，不点击进入正文。",
            "链接：只识别链接和聊天里附带的文字，不自动打开网页。",
            "视觉模型：可作为辅助，但需要严格校验，避免生成“对方分享了一段关于……”这类幻觉摘要。",
        ],
    )

    add_h1(doc, "九、制作步骤")
    add_numbers(
        doc,
        [
            "搭建基础桌面程序：选择 Python + Tkinter 或同类桌面 GUI，先实现主窗口、设置页、候选回复列表和发送按钮。",
            "接入窗口发现：使用 pywin32 读取前台窗口、窗口标题、进程名、窗口坐标和缩放信息。",
            "实现聊天区域截图：根据不同聊天软件的布局特征动态裁剪，避免硬编码固定坐标。",
            "接入 OCR：优先使用本地 RapidOCR/ONNXRuntime；保留 PaddleOCR 或其他 OCR 作为备用方案。",
            "实现角色识别：根据文本框、头像、气泡边界和消息区域位置判断“对方/我”，再合并多行长消息。",
            "接入文本大模型：封装 DeepSeek、智谱、通义千问、豆包、Kimi、硅基流动、自定义 HTTP 等接口。",
            "实现话术库：支持多格式导入，解析为上下文和回复样例，保存在本地 memory/vector 文件中。",
            "实现回复引擎：按回复来源模式生成候选，加入低质量回复过滤、复读过滤和编造数据拦截。",
            "实现发送器：聚焦目标聊天窗口，保护剪贴板，粘贴回复，按 Enter，并在发送前再次确认目标窗口。",
            "实现全托管：加入未回复轮次判断、已处理指纹、冷却时间、发送后复检和失败保护。",
            "打包交付：用 PyInstaller onedir 模式打包，包含 Python 运行时、OCR、OpenCV、ONNXRuntime、Excel/Word 解析依赖和图标。",
            "自检与发布：提供便携包自检、接口检测、OCR 排查说明，确保不打包 API Key、settings、memory 或训练数据。",
        ],
    )

    add_h1(doc, "十、测试清单")
    add_table(
        doc,
        ["测试项", "通过标准"],
        [
            ["窗口尺寸", "全屏、半屏、拖拽大小后仍能识别聊天区域和双方消息。"],
            ["窗口切换", "前台不是聊天窗口时不误识别；发送前目标窗口变化会拦截。"],
            ["OCR 噪声", "语音时长、时间戳、按钮、图片参数和乱码不会进入上下文。"],
            ["图片消息", "只有图片时不会生成虚构上下文；可见文字可以被识别。"],
            ["回复来源", "大模型模式不掺话术库；话术库模式可检索本地数据；混合模式可结合两者。"],
            ["导入文件", "txt/docx/xlsx/xls/zip 等可正确解析为话术库样例。"],
            ["全托管", "一轮只回一条；对方新增消息后才进入下一轮。"],
            ["便携包", "新电脑解压后填 API Key 即可用；不需要安装 Python 或 OCR 依赖。"],
        ],
        [4.0, 11.0],
    )

    add_h1(doc, "十一、当前方案的技术边界")
    add_bullets(
        doc,
        [
            "OCR 只能看到屏幕上可见内容，不能天然读取未滚动出来的历史消息。",
            "截图方案不调用聊天软件接口，因此稳定性取决于窗口布局、缩放比例、主题皮肤和遮挡情况。",
            "图片理解可以辅助，但视觉模型会幻觉，必须以可见文字和结构化校验为准。",
            "发送是键鼠自动化，不是官方接口；必须持续做目标窗口校验，降低误发风险。",
            "话术库增强不是模型微调，不能指望上传数据后模型权重发生变化。",
            "全托管适合低风险场景，高风险业务仍建议人工确认。",
        ],
    )

    add_h1(doc, "十二、后续可优化方向")
    add_table(
        doc,
        ["方向", "优化内容", "价值"],
        [
            ["窗口泛化", "增加更多聊天软件的窗口画像和区域定位规则", "支持更多应用，不局限微信。"],
            ["图像算法", "基于头像邻近、气泡轮廓和消息块检测做角色判断", "减少长文本跨中线导致的身份误判。"],
            ["向量检索", "把导入语料和手动确认回复写入本地向量库", "让话术库从关键词匹配升级到语义检索。"],
            ["视觉理解", "增加可见内容 OCR 与视觉模型双重校验", "看懂图片标题，但避免脑补。"],
            ["回复评估", "生成后自动评估是否回答最新问题、是否编造数据、是否复读", "提升候选回复质量。"],
            ["多联系人状态", "为每个聊天窗口维护独立上下文、已处理轮次和冷却状态", "避免不同会话串话。"],
            ["便携诊断", "把 OCR、API、依赖、权限、窗口识别状态做成一键诊断页", "降低客户反馈时的排查成本。"],
            ["安全策略", "全托管增加白名单、关键词阻断和高风险人工确认", "降低误发、骚扰和业务风险。"],
        ],
        [3.0, 7.2, 4.8],
    )

    add_h1(doc, "十三、交付建议")
    add_bullets(
        doc,
        [
            "发给客户时不要只发 exe，应发送完整 zip 或完整文件夹。",
            "便携包内不要包含打包者本机的 API Key、settings.json、memory.json、vector_memory.json 或 user_data。",
            "首次使用说明要写清楚：选择模型公司，填自己的 API Key，保存后再识别聊天。",
            "客户反馈问题时，优先收集：截图、设置页模型公司、OCR 上下文、回复来源、是否全托管、便携包自检报告。",
            "每次修复 OCR、角色判断、发送逻辑后，都应跑单元测试和 compileall，并在真实窗口里做一次手动识别验证。",
        ],
    )

    add_callout(
        doc,
        "一句话总结",
        "聊小智的核心不是“接入微信”，而是把人看屏幕、理解上下文、写回复、粘贴发送这套动作工程化。稳定性的关键在于窗口定位、OCR 清洗、角色判断、回复来源隔离和发送前校验。",
        LIGHT_GREEN,
    )

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    build()

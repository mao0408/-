from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("聊小智_制作思路与使用说明.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width_twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_twips))
    tc_w.set(qn("w:type"), "dxa")


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(31, 78, 121)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(subtitle)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(90, 90, 90)


def add_h1(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1"]
    p.add_run(text)


def add_h2(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 2"]
    p.add_run(text)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        set_cell_shading(hdr[i], "E8EEF5")
        set_cell_width(hdr[i], widths[i])
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
            set_cell_width(cells[i], widths[i])
    doc.add_paragraph()


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)

    for style_name, size, color in [
        ("Heading 1", 16, RGBColor(31, 78, 121)),
        ("Heading 2", 13, RGBColor(47, 84, 150)),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def main() -> None:
    doc = Document()
    configure_doc(doc)
    add_title(doc, "聊小智制作思路与使用说明", "Windows 端微信 OCR + 大模型回复助手的设计、运行逻辑与交付说明")

    add_h1(doc, "一、项目目标")
    add_para(
        doc,
        "这个程序的目标是在不调用微信官方接口的前提下，作为一个可视化伴随工具贴在微信窗口旁边，帮助用户读取当前聊天上下文，生成多条更像真人的候选回复，并在用户确认后发送到微信。便携版要求客户解压后即可运行，不需要安装 Python、OCR 或其它编程依赖；首次使用只需要在设置里选择模型公司并填写自己的 API Key。",
    )

    add_h1(doc, "二、整体运行流程")
    add_numbered(
        doc,
        [
            "启动程序后检测微信窗口位置，并把助手窗口吸附到微信右侧。",
            "用户点进某个具体聊天，点击“识别当前聊天”，或打开自动/全托管模式后由程序定时检测当前聊天变化。",
            "程序截取微信聊天区域截图，用本地 OCR 识别可见文字，并按气泡位置区分“对方”和“我”。",
            "清理 OCR 噪声，例如时间、重复文字、图标文字、语音秒数、无意义残片等。",
            "整理最近上下文，判断当前需要回复的是最近一轮未回复的对方消息，而不是旧话题。",
            "根据用户选择的回复来源，走“大模型”“话术库”或“大模型+话术库”生成候选回复。",
            "用户可以选择候选、修改候选、手动输入，点击发送后程序通过剪贴板把文本粘贴到微信输入框并按 Enter。",
            "只有用户手动确认发送或手动输入发送的内容会写入长期记忆；全托管发送不写入记忆。",
        ],
    )

    add_h1(doc, "三、为什么不需要微信 API")
    add_para(
        doc,
        "本程序不读取微信数据库，也不调用微信官方接口。它的核心方式是“屏幕可见内容理解 + 本地 OCR + 剪贴板输入”：微信窗口里已经显示出来的聊天内容会被截图识别，生成回复后再像用户复制粘贴一样把文本放进输入框。这样能避免微信接口权限问题，也降低了部署门槛。",
    )
    add_bullets(
        doc,
        [
            "识别来源：当前屏幕上可见的微信聊天内容。",
            "生成来源：用户配置的文本生成模型接口，或本地长期记忆检索结果。",
            "发送方式：复制文本到剪贴板，聚焦微信输入框，粘贴并回车。",
            "限制：只能理解当前可见或 OCR 能识别到的内容，不等同于读取完整微信聊天数据库。",
        ],
    )

    add_h1(doc, "四、核心模块设计")
    add_table(
        doc,
        ["模块", "职责"],
        [
            ["wechat_window.py", "检测微信窗口、获取坐标、判断当前聊天窗口、控制助手吸附位置。"],
            ["capture_ocr.py", "截取聊天区域，调用 RapidOCR 识别文字，按位置还原聊天顺序和角色。"],
            ["visible_content.py", "理解可见的图片、链接、卡片、表情等消息类型，表情包通常不生成回复。"],
            ["reply_engine.py", "组织上下文、回复目标、长期记忆和模型提示词，生成 3 条候选回复。"],
            ["memory.py", "保存完整上下文和用户最终发送的回复，提供相似历史检索。"],
            ["chat_training.py", "导入 txt、docx、xlsx、xls、csv、zip 等训练语料，并写入长期记忆。"],
            ["llm_clients.py", "对接智谱、DeepSeek、通义千问、豆包、Kimi、硅基流动和自定义 HTTP 文本接口。"],
            ["float_ui.py", "主浮窗 UI、设置页、回复来源切换、候选回复、手动输入、发送按钮等。"],
            ["sender.py", "隐藏浮窗、聚焦微信、粘贴回复并发送。"],
            ["settings.py", "保存模型公司、API Key、回复来源、OCR 参数、窗口参数等设置。"],
            ["self_check.py", "便携包自检，检查 OCR、截图、剪贴板、表格导入、依赖是否可用。"],
        ],
        [2400, 6960],
    )

    add_h1(doc, "五、回复生成逻辑")
    add_para(doc, "回复质量的关键不只是模型能力，还包括上下文整理、回复目标判断、提示词和候选筛选。当前逻辑按以下优先级处理：")
    add_bullets(
        doc,
        [
            "整段上下文作为参考，但对方最后一轮未回复消息优先级最高。",
            "如果对方连续发了多条且我方还没回复，应把这一轮作为统一回复目标。",
            "如果前面的问题已经由我方回复过，则只围绕最后一句或最后一轮新问题生成。",
            "大模型模式下不混入话术库，避免话术库污染当前回复。",
            "话术库模式下优先检索相似历史；没有合适数据时用大模型兜底。",
            "混合模式下先找相似历史片段，再交给大模型理解当前语境生成。",
            "候选不足时应继续让大模型补齐，而不是用本地通用敷衍句硬补。",
        ],
    )

    add_h1(doc, "六、训练和长期记忆机制")
    add_para(
        doc,
        "这里的“训练”不是对大模型做真正微调，而是长期记忆增强：程序保存完整的聊天上下文和你最终认可的回复，下次遇到相似场景时检索出来作为参考。这样成本低、部署简单，也适合便携版。",
    )
    add_bullets(
        doc,
        [
            "导入训练语料时，推荐使用 A/B 格式：A 代表对方，B 代表你希望学习的回复。",
            "程序会保存 B 方回复前 5-8 轮上下文和 B 方真实回复。",
            "支持 zip、txt、md、csv、tsv、json、docx、xlsx、xls、html 等格式。",
            "用户点击候选“发”、修改后发送、或手动输入发送，才会写入长期记忆。",
            "全托管模式无论是否发送成功，都不改变长期记忆，避免自动回复污染话术库。",
            "便携版的设置和记忆保存在程序同级 user_data 目录，不随初始便携包预置。",
        ],
    )

    add_h1(doc, "七、自动与全托管逻辑")
    add_table(
        doc,
        ["模式", "行为", "适用场景"],
        [
            ["手动识别", "用户点击按钮后识别当前聊天并生成候选，不自动发送。", "最稳，适合日常使用和测试。"],
            ["自动识别", "检测当前聊天变化后自动 OCR 并生成候选，不自动发送。", "OCR 区域稳定后使用，提高效率。"],
            ["全托管", "识别、生成、审核后自动发送最合适的一条。", "风险更高，需要先确认识别和回复质量稳定。"],
        ],
        [1800, 4200, 3360],
    )
    add_para(
        doc,
        "全托管必须遵守“一轮一回”的限制：只处理最近一条我方消息之后的连续对方消息；发送一条后记录该轮指纹；如果 OCR 再次看到同一轮消息，不再重复发送；只有对方新增消息后才进入下一轮。",
    )

    add_h1(doc, "八、便携版打包思路")
    add_bullets(
        doc,
        [
            "使用 PyInstaller onedir 打包，保留 _internal 文件夹承载 Python 运行时、OCR、OpenCV、ONNXRuntime 等依赖。",
            "不要只发单独 exe，必须发送完整便携包或 zip。",
            "打包时不带打包者本机的 API Key、settings.json、memory.json、user_data 或话术库。",
            "客户首次打开后，在设置里选择模型公司并填写自己的 API Key 即可。",
            "Base URL 和模型名会根据模型公司自动填好，普通用户不用修改。",
            "提供“便携包自检.bat”和“接口检测.bat”，便于定位 OCR 或 API 配置问题。",
        ],
    )

    add_h1(doc, "九、客户使用步骤")
    add_numbered(
        doc,
        [
            "解压完整便携包，双击“启动.bat”。",
            "打开微信电脑版，点进一个具体聊天窗口。",
            "进入设置，选择模型公司，例如 DeepSeek 或 Zhipu，然后填写自己的 API Key。",
            "回到主界面点击“识别当前聊天”，查看 OCR 上下文和候选回复。",
            "选择合适候选，点击“发”；不满意可以点击“改”或在下方手动输入。",
            "如需学习个人话术，在设置里导入自己的聊天记录训练包。",
            "确认识别稳定后，再考虑打开“自动”或“全托管”。",
        ],
    )

    add_h1(doc, "十、常见问题和注意事项")
    add_bullets(
        doc,
        [
            "OCR 识别不准：放大微信窗口，确保聊天区域无遮挡，避免窗口被其它软件覆盖。",
            "回复像兜底句：检查 API Key 是否生效、回复来源是否选对、模型返回是否为空或被过滤。",
            "话术库污染大模型模式：大模型模式应只走模型接口，不应拼入长期记忆。",
            "全托管连续回复：检查是否正确记录已处理轮次，并在发送后等待对方新增消息。",
            "换电脑无法 OCR：必须完整解压便携包，不能只复制 exe；_internal 文件夹不能删除。",
            "导入训练失败：确认文件格式是否受支持，Excel 老格式 xls 需要 xlrd 依赖，便携包已内置。",
        ],
    )

    add_h1(doc, "十一、后续优化方向")
    add_bullets(
        doc,
        [
            "继续优化 OCR 区域定位，减少误识别其它窗口或聊天列表内容。",
            "增加更强的可见内容理解，对图片、卡片、链接标题进行更稳的类型判断。",
            "在全托管前加入更严格的语境审核，必要时不发送。",
            "增加回复质量评分，过滤敷衍、复读、答非所问、过长或过营销的候选。",
            "话术库可加入更好的相似度算法，例如本地向量检索，而不是简单文本相似。",
            "设置页可继续简化，让普通用户只看到模型公司和 API Key，高级配置折叠起来。",
        ],
    )

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("版本说明：本文件为无二维码、无引流文案版本。")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.save(OUT)
    print(OUT.resolve())


if __name__ == "__main__":
    main()
